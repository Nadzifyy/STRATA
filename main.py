import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from groq import Groq
from pydantic import BaseModel, Field


load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / "static"

app = FastAPI(title="SWOT & TOWS Strategy Generator")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


class TowsRequest(BaseModel):
    strengths: list[str] = Field(default_factory=list)
    weaknesses: list[str] = Field(default_factory=list)
    opportunities: list[str] = Field(default_factory=list)
    threats: list[str] = Field(default_factory=list)


class RefineStrategyRequest(TowsRequest):
    strategy_type: Literal["so_strategies", "st_strategies", "wo_strategies", "wt_strategies"]
    current_strategy: str = Field(min_length=1, max_length=2000)
    instruction: str = Field(min_length=1, max_length=600)


class SwotAnalysis(BaseModel):
    strengths: list[str] = Field(default_factory=list)
    weaknesses: list[str] = Field(default_factory=list)
    opportunities: list[str] = Field(default_factory=list)
    threats: list[str] = Field(default_factory=list)


class ExportRequest(BaseModel):
    factors: TowsRequest = Field(default_factory=TowsRequest)
    swot_analysis: SwotAnalysis
    so_strategies: list[str] = Field(default_factory=list)
    st_strategies: list[str] = Field(default_factory=list)
    wo_strategies: list[str] = Field(default_factory=list)
    wt_strategies: list[str] = Field(default_factory=list)


SWOT_SECTIONS = (
    ("strengths", "Strengths", "Internal factors that support objectives"),
    ("weaknesses", "Weaknesses", "Internal factors that constrain objectives"),
    ("opportunities", "Opportunities", "External conditions that may be leveraged"),
    ("threats", "Threats", "External conditions that may impede progress"),
)

TOWS_SECTIONS = (
    ("so_strategies", "SO Strategies", "Leverage strengths to capture opportunities"),
    ("st_strategies", "ST Strategies", "Leverage strengths to mitigate threats"),
    ("wo_strategies", "WO Strategies", "Address weaknesses by pursuing opportunities"),
    ("wt_strategies", "WT Strategies", "Minimize weaknesses and avoid threats"),
)

FORMAL_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
FORMAL_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
FORMAL_BODY_FILL = "FFFFFF"


def pdf_safe(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u00d7": "x",
        "\u2192": "->",
        "\u2022": "-",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return re.sub(r"[^\x09\x0a\x0d\x20-\x7e]", "?", text)


def clean_list(items: list[str] | None) -> list[str]:
    return [str(item).strip() for item in (items or []) if item and str(item).strip()]


def format_cell_text(title: str, subtitle: str, items: list[str], empty: str = "None recorded.") -> str:
    lines = [title.upper(), subtitle, ""]
    values = items or [empty]
    lines.extend(f"{index}. {item}" for index, item in enumerate(values, 1))
    return "\n".join(lines)


def set_docx_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def set_docx_cell_borders(cell, color: str = "1A1A2E", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    tc_pr.append(borders)


def set_docx_cell_margins(cell, twips: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    margins = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:w"), str(twips))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    tc_pr.append(margins)


def fill_docx_cell(cell, title: str, subtitle: str, items: list[str]) -> None:
    cell.text = ""
    set_docx_cell_shading(cell, FORMAL_BODY_FILL)
    set_docx_cell_borders(cell)
    set_docx_cell_margins(cell)

    header = cell.paragraphs[0]
    header.paragraph_format.space_after = Pt(2)
    header_run = header.add_run(title.upper())
    header_run.bold = True
    header_run.font.name = "Times New Roman"
    header_run.font.size = Pt(11)
    header_run.font.color.rgb = FORMAL_NAVY

    subtitle_para = cell.add_paragraph()
    subtitle_para.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.italic = True
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(9)
    subtitle_run.font.color.rgb = FORMAL_GRAY

    values = items or ["None recorded."]
    for index, item in enumerate(values, 1):
        item_para = cell.add_paragraph()
        item_para.paragraph_format.space_after = Pt(4)
        item_para.paragraph_format.left_indent = Inches(0.05)
        item_run = item_para.add_run(f"{index}. {item}")
        item_run.font.name = "Times New Roman"
        item_run.font.size = Pt(10)
        item_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def add_docx_matrix(document: Document, sections: tuple, data: dict[str, list[str]]) -> None:
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, (key, label, subtitle) in enumerate(sections):
        row, col = divmod(index, 2)
        fill_docx_cell(table.cell(row, col), label, subtitle, clean_list(data.get(key)))
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.15)


def style_docx_heading(paragraph, size: int = 14) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.color.rgb = FORMAL_NAVY
        run.font.all_caps = False


def build_docx(payload: ExportRequest) -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    for section in document.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("STRATEGIC ANALYSIS REPORT")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = FORMAL_NAVY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("SWOT Assessment and TOWS Strategy Matrix")
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = FORMAL_GRAY

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    meta_run = meta.add_run(f"Prepared by Strata  |  {datetime.now().strftime('%d %B %Y')}")
    meta_run.font.name = "Times New Roman"
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = FORMAL_GRAY

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_run = rule.add_run("────────────────────────────────────────")
    rule_run.font.color.rgb = FORMAL_GRAY
    rule_run.font.size = Pt(10)

    intro = document.add_paragraph(
        "This report presents a structured SWOT analysis and corresponding TOWS strategies "
        "derived exclusively from the factors supplied by the user. The matrices below may be "
        "edited for further planning and decision-making."
    )
    intro.paragraph_format.space_after = Pt(12)
    for run in intro.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    factors = {
        "strengths": clean_factors(payload.factors.strengths),
        "weaknesses": clean_factors(payload.factors.weaknesses),
        "opportunities": clean_factors(payload.factors.opportunities),
        "threats": clean_factors(payload.factors.threats),
    }
    section_number = 1
    if any(factors.values()):
        heading = document.add_paragraph()
        heading_run = heading.add_run(f"{section_number}. Input Factors")
        style_docx_heading(heading)
        tagged = {
            key: [f"[{key[0].upper()}{index}] {item}" for index, item in enumerate(items, 1)]
            for key, items in factors.items()
        }
        add_docx_matrix(document, SWOT_SECTIONS, tagged)
        section_number += 1

    heading = document.add_paragraph()
    heading_run = heading.add_run(f"{section_number}. SWOT Analysis")
    style_docx_heading(heading)
    add_docx_matrix(document, SWOT_SECTIONS, payload.swot_analysis.model_dump())
    section_number += 1

    heading = document.add_paragraph()
    heading_run = heading.add_run(f"{section_number}. TOWS Strategy Matrix")
    style_docx_heading(heading)
    add_docx_matrix(document, TOWS_SECTIONS, payload.model_dump())

    footer = document.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    footer_run = footer.add_run(
        "Confidential — Intended for internal planning use. Strategies should be validated against operational constraints before implementation."
    )
    footer_run.italic = True
    footer_run.font.name = "Times New Roman"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = FORMAL_GRAY

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


class FormalReportPDF(FPDF):
    def header(self) -> None:
        self.set_font("Times", "B", 9)
        self.set_text_color(26, 26, 46)
        self.cell(0, 6, "STRATA  |  STRATEGIC ANALYSIS REPORT", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_draw_color(26, 26, 46)
        self.set_line_width(0.4)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(6)

    def footer(self) -> None:
        self.set_y(-14)
        self.set_draw_color(26, 26, 46)
        self.set_line_width(0.3)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(2)
        self.set_font("Times", "I", 8)
        self.set_text_color(90, 90, 90)
        self.cell(0, 5, f"Page {self.page_no()}  |  Confidential", align="C")


def build_pdf(payload: ExportRequest) -> bytes:
    pdf = FormalReportPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    def write(text: str, size: int = 11, style: str = "", gap: float = 1, align: str = "L") -> None:
        pdf.set_font("Times", style, size)
        pdf.set_text_color(34, 34, 34)
        pdf.multi_cell(
            0,
            max(5, size * 0.5),
            pdf_safe(text),
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
            align=align,
        )
        if gap:
            pdf.ln(gap)

    def add_matrix(sections: tuple, data: dict[str, list[str]]) -> None:
        pdf.set_font("Times", "", 9)
        pdf.set_text_color(34, 34, 34)
        with pdf.table(
            col_widths=(pdf.epw / 2, pdf.epw / 2),
            text_align="LEFT",
            line_height=4.2,
            padding=4,
            borders_layout="ALL",
        ) as table:
            for row_index in range(2):
                row = table.row()
                for col_index in range(2):
                    key, label, subtitle = sections[row_index * 2 + col_index]
                    row.cell(
                        pdf_safe(
                            format_cell_text(
                                label,
                                subtitle,
                                clean_list(data.get(key)),
                            )
                        )
                    )
        pdf.ln(5)

    write("STRATEGIC ANALYSIS REPORT", size=16, style="B", gap=1, align="C")
    write("SWOT Assessment and TOWS Strategy Matrix", size=11, style="I", gap=1, align="C")
    write(f"Prepared by Strata  |  {datetime.now().strftime('%d %B %Y')}", size=9, gap=3, align="C")
    write(
        "This report presents a structured SWOT analysis and corresponding TOWS strategies "
        "derived exclusively from the factors supplied by the user.",
        size=10,
        gap=5,
    )

    factors = {
        "strengths": clean_factors(payload.factors.strengths),
        "weaknesses": clean_factors(payload.factors.weaknesses),
        "opportunities": clean_factors(payload.factors.opportunities),
        "threats": clean_factors(payload.factors.threats),
    }
    section_number = 1
    if any(factors.values()):
        write(f"{section_number}. Input Factors", size=12, style="B", gap=2)
        tagged = {
            key: [f"[{key[0].upper()}{index}] {item}" for index, item in enumerate(items, 1)]
            for key, items in factors.items()
        }
        add_matrix(SWOT_SECTIONS, tagged)
        section_number += 1

    write(f"{section_number}. SWOT Analysis", size=12, style="B", gap=2)
    add_matrix(SWOT_SECTIONS, payload.swot_analysis.model_dump())
    section_number += 1

    write(f"{section_number}. TOWS Strategy Matrix", size=12, style="B", gap=2)
    add_matrix(TOWS_SECTIONS, payload.model_dump())

    write(
        "Confidential - Intended for internal planning use. Strategies should be validated "
        "against operational constraints before implementation.",
        size=8,
        style="I",
        gap=0,
    )

    return bytes(pdf.output())


SYSTEM_PROMPT = (
    "You are a pure SWOT and TOWS Strategic Engine. ONLY use provided factors. "
    "Do NOT add outside knowledge. Return strictly valid JSON containing "
    "swot_analysis, so_strategies, st_strategies, wo_strategies, and "
    "wt_strategies. swot_analysis must be an object with strengths, weaknesses, "
    "opportunities, and threats arrays. Every analysis item and strategy must use "
    "factor citations like [S1], [O1]."
)


def clean_factors(items: list[str]) -> list[str]:
    return [item.strip() for item in items if item and item.strip()]


def format_factors(factors: dict[str, list[str]]) -> str:
    return "\n".join(
        f"{label}:\n" + "\n".join(f"[{tag}{index}] {factor}" for index, factor in enumerate(factors[key], 1))
        for label, tag, key in (
            ("Strengths", "S", "strengths"),
            ("Weaknesses", "W", "weaknesses"),
            ("Opportunities", "O", "opportunities"),
            ("Threats", "T", "threats"),
        )
    )


@app.get("/", include_in_schema=False)
async def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.post("/api/generate-tows")
async def generate_tows(payload: TowsRequest) -> dict:
    factors = {
        "strengths": clean_factors(payload.strengths),
        "weaknesses": clean_factors(payload.weaknesses),
        "opportunities": clean_factors(payload.opportunities),
        "threats": clean_factors(payload.threats),
    }

    if not all(factors.values()):
        raise HTTPException(
            status_code=400,
            detail="Add at least one factor to every S, W, O, and T category.",
        )

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key or api_key == "your_groq_api_key_here":
        raise HTTPException(
            status_code=500,
            detail="Set a valid GROQ_API_KEY in your .env file before generating strategies.",
        )

    formatted_factors = format_factors(factors)

    user_prompt = (
        "First, create a concise SWOT analysis: interpret each category using only "
        "the supplied factors. Then create concise, actionable TOWS strategies. "
        "Return JSON in this exact shape: {\"swot_analysis\": {\"strengths\": [], "
        "\"weaknesses\": [], \"opportunities\": [], \"threats\": []}, "
        "\"so_strategies\": [], \"st_strategies\": [], \"wo_strategies\": [], "
        "\"wt_strategies\": []}.\n\n"
        f"{formatted_factors}"
    )

    try:
        client = Groq(api_key=api_key)
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.25,
        )
        result = json.loads(completion.choices[0].message.content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail="Groq returned invalid JSON.") from exc
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Unable to generate strategies: {exc}") from exc

    required_keys = ("so_strategies", "st_strategies", "wo_strategies", "wt_strategies")
    swot_keys = ("strengths", "weaknesses", "opportunities", "threats")
    if not all(isinstance(result.get(key), list) for key in required_keys):
        raise HTTPException(status_code=502, detail="Groq returned an unexpected strategy format.")
    if not isinstance(result.get("swot_analysis"), dict) or not all(
        isinstance(result["swot_analysis"].get(key), list) for key in swot_keys
    ):
        raise HTTPException(status_code=502, detail="Groq returned an unexpected SWOT format.")

    return {
        "swot_analysis": {key: result["swot_analysis"][key] for key in swot_keys},
        **{key: result[key] for key in required_keys},
    }


@app.post("/api/refine-strategy")
async def refine_strategy(payload: RefineStrategyRequest) -> dict:
    factors = {
        "strengths": clean_factors(payload.strengths),
        "weaknesses": clean_factors(payload.weaknesses),
        "opportunities": clean_factors(payload.opportunities),
        "threats": clean_factors(payload.threats),
    }
    if not all(factors.values()):
        raise HTTPException(status_code=400, detail="Add at least one factor to every S, W, O, and T category.")

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key or api_key == "your_groq_api_key_here":
        raise HTTPException(status_code=500, detail="Set a valid GROQ_API_KEY in your environment before refining strategies.")

    strategy_label = payload.strategy_type.replace("_strategies", "").upper()
    user_prompt = (
        f"Refine this {strategy_label} strategy using the user's direction. Return exactly "
        "one concise strategy, preserve or improve its relevant factor citations, and use "
        "only the factors below. Return JSON: {\"refined_strategy\": \"...\"}.\n\n"
        f"Current strategy: {payload.current_strategy}\n"
        f"User direction: {payload.instruction.strip()}\n\n"
        f"{format_factors(factors)}"
    )

    try:
        client = Groq(api_key=api_key)
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": "You are a pure TOWS Strategic Engine. ONLY use provided factors. Do NOT add outside knowledge. Return strictly valid JSON containing refined_strategy with factor citations like [S1], [O1].",
                },
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.25,
        )
        result = json.loads(completion.choices[0].message.content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail="Groq returned invalid JSON.") from exc
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Unable to refine strategy: {exc}") from exc

    refined_strategy = result.get("refined_strategy")
    if not isinstance(refined_strategy, str) or not refined_strategy.strip():
        raise HTTPException(status_code=502, detail="Groq returned an unexpected refinement format.")

    return {"refined_strategy": refined_strategy.strip()}


@app.post("/api/export-docx")
async def export_docx(payload: ExportRequest) -> StreamingResponse:
    try:
        content = build_docx(payload)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Unable to build Word document: {exc}") from exc
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="strata-swot-tows.docx"'},
    )


@app.post("/api/export-pdf")
async def export_pdf(payload: ExportRequest) -> StreamingResponse:
    try:
        content = build_pdf(payload)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Unable to build PDF: {exc}") from exc
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="strata-swot-tows.pdf"'},
    )
